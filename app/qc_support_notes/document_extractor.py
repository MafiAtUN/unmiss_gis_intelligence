"""Document extraction utilities for Word and PDF files."""

from typing import Optional
from pathlib import Path
import io


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a Word document (.docx).
    
    Args:
        file_path: Path to the .docx file.
        
    Returns:
        Extracted text or None if error.
    """
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        return None


def extract_text_from_docx_bytes(file_bytes: bytes) -> Optional[str]:
    """
    Extract text from Word document bytes.
    
    Args:
        file_bytes: Bytes of the .docx file.
        
    Returns:
        Extracted text or None if error.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        return None


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file.
        
    Returns:
        Extracted text or None if error.
    """
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            return "\n".join(text_parts)
    except ImportError:
        # Try pdfplumber as fallback
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text_parts.append(page.extract_text())
                return "\n".join(text_parts)
        except ImportError:
            return None
    except Exception as e:
        return None


def extract_text_from_pdf_bytes(file_bytes: bytes) -> Optional[str]:
    """
    Extract text from PDF bytes.
    
    Args:
        file_bytes: Bytes of the PDF file.
        
    Returns:
        Extracted text or None if error.
    """
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in pdf_reader.pages:
            text_parts.append(page.extract_text())
        return "\n".join(text_parts)
    except ImportError:
        # Try pdfplumber as fallback
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text_parts.append(page.extract_text())
                return "\n".join(text_parts)
        except ImportError:
            return None
    except Exception as e:
        return None


def extract_text_from_file(file_path: Optional[str] = None, file_bytes: Optional[bytes] = None, 
                           file_name: Optional[str] = None) -> Optional[str]:
    """
    Extract text from a file (Word or PDF) based on extension.
    
    Args:
        file_path: Path to the file.
        file_bytes: Bytes of the file (for uploaded files).
        file_name: Name of the file (to determine type).
        
    Returns:
        Extracted text or None if error.
    """
    # Determine file type
    if file_name:
        file_ext = Path(file_name).suffix.lower()
    elif file_path:
        file_ext = Path(file_path).suffix.lower()
    else:
        return None
    
    if file_ext == '.docx':
        if file_bytes:
            return extract_text_from_docx_bytes(file_bytes)
        elif file_path:
            return extract_text_from_docx(file_path)
    elif file_ext == '.pdf':
        if file_bytes:
            return extract_text_from_pdf_bytes(file_bytes)
        elif file_path:
            return extract_text_from_pdf(file_path)
    
    return None

