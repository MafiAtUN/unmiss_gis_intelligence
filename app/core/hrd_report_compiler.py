"""Compile HRD Daily Reports from field office dailies."""
import re
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.hrd_incident_extractor import HRDIncidentExtractor, HRDIncident
from app.utils.logging import log_error


class HRDReportCompiler:
    """Compile HRD Daily Reports from field office dailies."""
    
    def __init__(self, incident_extractor: Optional[HRDIncidentExtractor] = None):
        """
        Initialize report compiler.
        
        Args:
            incident_extractor: HRDIncidentExtractor instance
        """
        self.incident_extractor = incident_extractor or HRDIncidentExtractor()
    
    def compile_daily_report(self, 
                           date: datetime,
                           field_office_dailies: List[Dict[str, Any]],
                           output_path: str) -> str:
        """
        Compile HRD Daily Report from field office dailies.
        
        Args:
            date: Report date
            field_office_dailies: List of dicts with 'file_path' and 'field_office'
            output_path: Path to save compiled report
            
        Returns:
            Path to compiled report
        """
        # Extract incidents from all dailies
        all_incidents = []
        for daily_info in field_office_dailies:
            file_path = daily_info.get("file_path")
            field_office = daily_info.get("field_office")
            
            if not file_path:
                continue
            
            # Read document
            try:
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                
                # Extract incidents
                incidents = self.incident_extractor.extract_incidents_from_text(text, field_office)
                all_incidents.extend(incidents)
            except Exception as e:
                log_error(e, {
                    "module": "hrd_report_compiler",
                    "file_path": file_path,
                    "field_office": field_office
                })
        
        # Generate compiled report
        doc = self._create_compiled_report(date, all_incidents)
        doc.save(output_path)
        
        return output_path
    
    def _create_compiled_report(self, date: datetime, incidents: List[HRDIncident]) -> Document:
        """Create compiled HRD Daily Report document."""
        doc = Document()
        
        # Header
        header_para = doc.add_paragraph("UNITED NATIONS         ألأمم المتحدة")
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("United Nations Mission in South Sudan (UNMISS)")
        doc.add_paragraph("Human Rights Division (HRD)")
        doc.add_paragraph("Daily Situation Report")
        doc.add_paragraph(date.strftime("%d %B %Y"))
        
        # Highlights
        doc.add_paragraph("Highlights")
        highlights = self._generate_highlights(incidents)
        for highlight in highlights:
            doc.add_paragraph(highlight)
        
        # Group incidents by state
        incidents_by_state = self._group_by_state(incidents)
        
        # Detailed incidents by state
        for state, state_incidents in incidents_by_state.items():
            doc.add_paragraph(state)
            doc.add_paragraph()  # Empty line
            
            for incident in state_incidents:
                # Add incident description
                if incident.description:
                    doc.add_paragraph(incident.description)
                
                # Add comments if any (from description analysis)
                comments = self._extract_comments(incident.description)
                if comments:
                    doc.add_paragraph(f"Comments: {comments}")
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("End    -")
        
        return doc
    
    def _generate_highlights(self, incidents: List[HRDIncident]) -> List[str]:
        """Generate highlights section from incidents."""
        highlights = []
        
        for incident in incidents[:10]:  # Top 10 incidents
            if not incident.description:
                continue
            
            # Create brief highlight
            violation = incident.types_of_violations or "incident"
            location = incident.location_of_incident or "unknown location"
            state = incident.incident_state or ""
            
            # Extract key info for highlight
            highlight = f"{incident.alleged_perpetrators or 'Unknown perpetrators'} {violation.lower()} "
            
            if incident.total_victims:
                if incident.total_victims == 1:
                    highlight += "one civilian"
                else:
                    highlight += f"{incident.total_victims} civilians"
            else:
                highlight += "civilians"
            
            highlight += f" in {location}"
            if state:
                highlight += f" ({state})"
            highlight += ";"
            
            highlights.append(highlight)
        
        return highlights
    
    def _group_by_state(self, incidents: List[HRDIncident]) -> Dict[str, List[HRDIncident]]:
        """Group incidents by state."""
        grouped = {}
        for incident in incidents:
            state = incident.incident_state or "Unknown"
            if state not in grouped:
                grouped[state] = []
            grouped[state].append(incident)
        return grouped
    
    def _extract_comments(self, description: Optional[str]) -> Optional[str]:
        """Extract comments from description."""
        if not description:
            return None
        
        # Look for "Comment:" or "Comments:" patterns
        comment_pattern = r'(?:Comment|Comments):\s*(.+?)(?:\n|$)'
        match = re.search(comment_pattern, description, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        
        return None

