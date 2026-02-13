#!/usr/bin/env python3
"""
Evaluate HRD extraction performance by comparing generated reports with originals.
"""

import sys
from pathlib import Path
from docx import Document
import pandas as pd
from typing import Dict, List, Any
import difflib

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))


def compare_compiled_reports(generated_path: str, original_path: str) -> Dict[str, Any]:
    """Compare generated compiled report with original."""
    try:
        gen_doc = Document(generated_path)
        orig_doc = Document(original_path)
        
        gen_text = [para.text.strip() for para in gen_doc.paragraphs if para.text.strip()]
        orig_text = [para.text.strip() for para in orig_doc.paragraphs if para.text.strip()]
        
        # Count incidents in each
        gen_incidents = sum(1 for text in gen_text if any(keyword in text.lower() for keyword in ['reported', 'sources', 'killed', 'injured', 'arrested']))
        orig_incidents = sum(1 for text in orig_text if any(keyword in text.lower() for keyword in ['reported', 'sources', 'killed', 'injured', 'arrested']))
        
        # Calculate similarity
        gen_full = "\n".join(gen_text)
        orig_full = "\n".join(orig_text)
        similarity = difflib.SequenceMatcher(None, gen_full.lower(), orig_full.lower()).ratio()
        
        return {
            "generated_paragraphs": len(gen_text),
            "original_paragraphs": len(orig_text),
            "generated_incidents": gen_incidents,
            "original_incidents": orig_incidents,
            "similarity": similarity,
            "generated_text": gen_text,
            "original_text": orig_text
        }
    except Exception as e:
        return {"error": str(e)}


def compare_matrices(generated_path: str, original_path: str) -> Dict[str, Any]:
    """Compare generated matrix with original."""
    try:
        gen_df = pd.read_excel(generated_path)
        orig_df = pd.read_excel(original_path)
        
        # Compare structure
        gen_cols = set(gen_df.columns)
        orig_cols = set(orig_df.columns)
        
        # Compare key fields
        comparison = {
            "generated_rows": len(gen_df),
            "original_rows": len(orig_df),
            "generated_columns": len(gen_df.columns),
            "original_columns": len(orig_df.columns),
            "common_columns": len(gen_cols & orig_cols),
            "missing_columns": list(orig_cols - gen_cols),
            "extra_columns": list(gen_cols - orig_cols),
        }
        
        # Sample comparison
        if len(gen_df) > 0 and len(orig_df) > 0:
            gen_sample = gen_df.iloc[0]
            orig_sample = orig_df.iloc[0] if len(orig_df) > 0 else None
            
            comparison["sample_generated"] = {
                "incident_code": gen_sample.get("Incident Code"),
                "date_incident": str(gen_sample.get("Date of Incident")),
                "reporting_fo": gen_sample.get("Reporting Field Office"),
                "location": gen_sample.get("Location of Incident"),
                "violation": gen_sample.get("Types of violations"),
            }
            
            if orig_sample is not None:
                comparison["sample_original"] = {
                    "incident_code": orig_sample.get("Incident Code"),
                    "date_incident": str(orig_sample.get("Date of Incident")),
                    "reporting_fo": orig_sample.get("Reporting Field Office"),
                    "location": orig_sample.get("Location of Incident"),
                    "violation": orig_sample.get("Types of violations"),
                }
        
        return comparison
    except Exception as e:
        return {"error": str(e)}


def main():
    """Main evaluation function."""
    week_folder = Path("resources/Weekly/03-09")
    date = "4 November 2025"
    
    generated_report = week_folder / f"HRD Daily Report_{date}.docx"
    original_report = week_folder / f"HRD Daily Report_{date}.docx"  # Same file for now
    
    generated_matrix = week_folder / "Weekly CivCas Matrix-03-03 November 2025.xlsx"
    original_matrix = week_folder / "Weekly CivCas Matrix-3-9 November 2025.xlsx"
    
    print("=" * 80)
    print("HRD EXTRACTION EVALUATION")
    print("=" * 80)
    
    # Check if generated report exists
    if generated_report.exists():
        print(f"\n✓ Generated report found: {generated_report.name}")
        
        # Read and display generated report
        doc = Document(generated_report)
        print(f"\n=== GENERATED COMPILED REPORT ({date}) ===")
        print("-" * 80)
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                print(f"{i:2d}. {text}")
                if i > 50:  # Limit output
                    print("... (truncated)")
                    break
    else:
        print(f"\n✗ Generated report not found: {generated_report}")
    
    # Compare matrices
    if generated_matrix.exists() and original_matrix.exists():
        print("\n" + "=" * 80)
        print("MATRIX COMPARISON")
        print("=" * 80)
        
        comparison = compare_matrices(str(generated_matrix), str(original_matrix))
        
        if "error" not in comparison:
            print(f"\nGenerated Matrix:")
            print(f"  Rows: {comparison['generated_rows']}")
            print(f"  Columns: {comparison['generated_columns']}")
            
            print(f"\nOriginal Matrix:")
            print(f"  Rows: {comparison['original_rows']}")
            print(f"  Columns: {comparison['original_columns']}")
            
            print(f"\nCommon columns: {comparison['common_columns']}")
            
            if "sample_generated" in comparison:
                print(f"\n=== SAMPLE FROM GENERATED MATRIX ===")
                sample = comparison["sample_generated"]
                for key, value in sample.items():
                    print(f"  {key}: {value}")
            
            if "sample_original" in comparison:
                print(f"\n=== SAMPLE FROM ORIGINAL MATRIX ===")
                sample = comparison["sample_original"]
                for key, value in sample.items():
                    print(f"  {key}: {value}")
        else:
            print(f"Error: {comparison['error']}")
    
    print("\n" + "=" * 80)
    print("EVALUATION COMPLETE")
    print("=" * 80)


if __name__ == "__main__":
    main()

